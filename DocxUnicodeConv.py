#! /usr/local/bin/python
# -*- coding: utf-8 -*-
#
# The concept behind this code is taken from page
# http://stackoverflow.com/questions/17850227/text-replace-in-docx-and-save-the-changed-file-with-python-docx

import os
import sys
import getopt
import xml.etree.ElementTree as ET

from docx import Document

from XmlExtract import extract_xml
from XmlExtract import create_newdoc
from BaminiDict import bamini_dict
from AmudhamDict import amudham_dict
from AdhawinTamilDict import adhawintamil_dict 
from TamilFancyDict import tamilfancy_dict 


# Global variables
english_fonts = ["Times New Roman"]


# Functions....
def usage():
    print ""
    print "This tool converts word document (.docx) file with Bamini font to"
    print "Unicode font without changing the format of the document"
    print ""
    print "usage: python DocxUnicodeConv.py -i <input file> [ -p <outpath> ]"
    print ""


def print_run(run):
    print "run font = "+str(run.font.name)
    print "run style.font = "+str(run.style.font.name)
    print "run text = "+run.text


def convert_amudham(wg):
    text = wg.encode("utf-8")
    for key in amudham_dict.keys():
        text = text.replace(str(key), str(amudham_dict.get(key)))
    print "Amudham: "+text
    return text.decode("utf-8")


def convert_bamini(wg):
    text = wg.encode("utf-8")
    for key in bamini_dict.keys():
        text = text.replace(str(key), str(bamini_dict.get(key)))
    print "Bamini: "+text
    return text.decode("utf-8")


def convert_tamilfancy(wg):
    text = wg.encode("utf-8")
    # dependent vowel here comes with 2 bytes
    igh = 'à'
    pos = text.find(igh)
    if pos != -1:
        # found! swap characters round the vowel to render it correctly in unicode
        ltext = list(text)
        temp = ltext[pos+2]
        ltext[pos+2] = ltext[pos-1]
        ltext[pos-1] = temp
        text = "".join(ltext)
        print text

    # real conversion
    for key in tamilfancy_dict.keys():
        text = text.replace(str(key), str(tamilfancy_dict.get(key)))

    print "Tamil_Fancy: "+text
    return text.decode("utf-8")


def convert_adhawintamil(wg):
    text = wg.encode("utf-8")
    for key in adhawintamil_dict.keys():
        text = text.replace(str(key), str(adhawintamil_dict.get(key)))
    print "Adhawin-Tamil: "+text
    return text.decode("utf-8")


# arg1: run object, arg2: paragraph font
def convert_runfont(run, p_font):
    print_run(run)
    r_font = run.font.name

    # ignore runs with text length 0 or less
    if len(run.text) <= 0:
        print " * * *  R U N   L E N G T H   Z E R O  * * *"
        return

    # main convversion logic
    if r_font == "Amudham" or (r_font == None and p_font == "Amudham"):
        run.text = convert_amudham(run.text)
    elif r_font == "Adhawin-Tamil" or (r_font == None and p_font == "Adhawin-Tamil"):
        run.text = convert_adhawintamil(run.text)
    elif r_font == "Bamini" or (r_font == None and p_font == "Bamini"):
        run.text = convert_bamini(run.text)
    elif r_font == "Tamil_Fancy" or (r_font == None and p_font == "Tamil_Fancy"):
        run.text = convert_tamilfancy(run.text)
    elif r_font in english_fonts:
        print "No conversion for "+str(r_font)
    else:
        print " * * *   U N K N O W N   F O N T   * * * "


def DocxUnicodeConv(infile, outpath):
    outfile = outpath+infile+".mod.docx"
    document = Document(infile)
    for p in document.paragraphs:
        # Collect paragraph information
        print "\n<para>"
        paragraph_font = str(p.style.font.name)
        print "paragraph font: "+paragraph_font
        print "paragraph text: "+p.text

        # Collect run information in this paragraph
        runs = [r for r in p.runs]
        runs_len = len(runs)
        print "No of runs: "+ str(runs_len)

        # if no runs, skip conversion
        if runs_len == 0:
            print "</para>"
            continue

        # Convert single-run paragragh
        if runs_len == 1:
            print "\n<run>"
            convert_runfont(runs[0], paragraph_font)
            print "</run>"
            print "</para>"
            continue

        # Convert multi-run paragraph
        i = ref = 0
        pending_i_vowel = None
        #for run in p.runs:
        for i in range(runs_len):
            print "\n<run> "+str(i)

            # Search text in run: ignore all zero lengthed texts
            if len(runs[i].text) <= 0:
                print "</run>"
                continue

            # take a copy of i as reference
            ref = i

            # Cat all run.text as long as they have same font
            while i+1 < runs_len and runs[i].font.name == runs[i+1].font.name:
                runs[ref].text += runs[i+1].text
                runs[i+1].clear() # clear the content of this run
                i += 1

            # Correct the ¿ ை(  )error by not coverting immediately
            if (runs[ref].font.name == "Adhawin-Tamil" or (paragraph_font == "Adhawin-Tamil" and runs[ref].font.name == None)) and runs[ref].text[0].encode("utf-8") == '¿':
                print "P E N D I N G   V O W E L   'I'   D E T E C T E D"
                pending_i_vowel = runs[ref].text
                runs[ref].clear()
            # if ¿ symbol was captured, then insert it in the 2nd index of the next string
            elif pending_i_vowel != None:
                runs[ref].text = runs[ref].text[:1]+pending_i_vowel+runs[ref].text[1:]
                convert_runfont(runs[ref], paragraph_font)
                print "P E N D I N G   V O W E L   'I'   C O R R E C T E D"
                pending_i_vowel = None
            else:
                convert_runfont(runs[ref], paragraph_font)

            print "</run>"

        print "</para>"
 

    # Now search for text in tables
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        if len(run.text) > 0:
                            convert_runfont(run, p.style.font.name)

    # Save it before the document is extracted for textbox
    modified_file = infile+"-modified.docx"
    document.save(modified_file)
    
    # Now search text in textboxes and replace them
    docxml = extract_xml(modified_file)
    newxml = ParseAndReplaceTextBoxTexts(docxml)
    create_newdoc(modified_file, newxml)

    # Clean-up the directory
    os.remove(modified_file)
    os.remove(docxml)
    os.remove(newxml)


def ParseAndReplaceTextBoxTexts(filepath):
    tree = ET.parse(filepath)
    root = tree.getroot()
    pending_tbi_vowel = None

    # Parse text boxes and replace
    for i in range(len(root[0])):
        for j in range(len(root[0][i])):
            for k in range(len(root[0][i][j])):
                for l in range(len(root[0][i][j][k])):
                    if root[0][i][j][k][l].tag.split('}')[-1] == 'Fallback':
                        for m in range(len(root[0][i][j][k][l])):
                            for n in range(len(root[0][i][j][k][l][m])):
                                for o in range(len(root[0][i][j][k][l][m][n])):
                                    for p in range(len(root[0][i][j][k][l][m][n][o])):
                                        if root[0][i][j][k][l][m][n][o][p].tag.split('}')[-1] == 'textbox':
                                            for q in range(len(root[0][i][j][k][l][m][n][o][p])):
                                                for r in range(len(root[0][i][j][k][l][m][n][o][p][q])):
                                                    for s in range(len(root[0][i][j][k][l][m][n][o][p][q][r])):
                                                        for t in range(len(root[0][i][j][k][l][m][n][o][p][q][r][s])):
                                                            if root[0][i][j][k][l][m][n][o][p][q][r][s][t].tag.split('}')[-1] == 't':
                                                                tb_text =  root[0][i][j][k][l][m][n][o][p][q][r][s][t].text
                                                                if tb_text.encode("utf-8") == '¿':
                                                                    pending_tbi_vowel = tb_text
                                                                    root[0][i][j][k][l][m][n][o][p][q][r][s][t].text = ""
                                                                elif pending_tbi_vowel != None:
                                                                    tb_text = tb_text[:1]+pending_tbi_vowel+tb_text[1:]
                                                                    root[0][i][j][k][l][m][n][o][p][q][r][s][t].text = convert_bamini(tb_text)
                                                                    pending_tbi_vowel = None
                                                                else:
                                                                    root[0][i][j][k][l][m][n][o][p][q][r][s][t].text = convert_bamini(tb_text)

    # Write the modified xml file
    modxml = filepath+"-modified.xml"
    tree.write(modxml)
    return modxml



def main():
    try:
        opts, args = getopt.getopt(sys.argv[1:], "i:p:", ["help", "input", "path"])
    except getopt.GetoptError as err:
        # print help information and exit:
        print str(err)  # will print something like "option -a not recognized"
        usage()
        sys.exit(2)
    
    infile = ""
    outfile = None
    outpath = "./"
    for o, a in opts:
        if o in ("-i", "--input"):
            infile = a
        elif o in ("-p", "--path"):
            outpath = a
        elif o in ("-h", "--help"):
            usage()
            sys.exit()
        else:
            print "option = "+o
            assert False, "unhandled option"
    
    if infile == "":
        print ""
        print "Error: Not all expected arguments are passed!!"
        print ""
        usage()
        sys.exit(2)
   
    if os.path.isfile(infile):
        DocxUnicodeConv(infile, outpath)
            


if __name__ == "__main__": 
    main()
